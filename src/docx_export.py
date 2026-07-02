"""Write an assembled proposal draft to a .docx file."""

from pathlib import Path

from docx import Document


def write_docx(sections: list[tuple[str, str]], out_path: str, title: str) -> None:
    """Write drafted sections to a formatted .docx file.

    Each section is a (question, answer) pair. The question becomes a heading
    and the drafted answer appears as paragraphs beneath it.
    """
    document = Document()
    document.add_heading(title, level=0)

    for question, answer in sections:
        document.add_heading(question, level=1)
        for paragraph in answer.split("\n\n"):
            cleaned = paragraph.strip()
            if cleaned:
                document.add_paragraph(cleaned)

    destination = Path(out_path)
    if destination.parent and not destination.parent.exists():
        destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))